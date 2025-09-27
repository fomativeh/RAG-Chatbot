import os
from typing import List, Tuple

import streamlit as st
from dotenv import load_dotenv
from openai import OpenAI
import logging

from langchain_openai import ChatOpenAI
from langchain.schema import SystemMessage, HumanMessage, Document as LCDocument
from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter

from utils import get_or_create_collection, get_or_create_session_collection, delete_session_collection, get_models, get_chunking, get_env
from prompts import build_system_prompt, format_context, build_user_prompt, build_safety_check_prompt
from metadata import infer_file_metadata, extract_chunk_section, citation_label


load_dotenv()
# --- Logging setup (console) ---
logger = logging.getLogger("rag")
if not logger.handlers:
    logger.setLevel(logging.INFO)
    sh = logging.StreamHandler()
    fmt = logging.Formatter("%(asctime)s [%(levelname)s] %(message)s")
    sh.setFormatter(fmt)
    logger.addHandler(sh)


DEFAULT_COLLECTION = "rag_docs"  # pre-ingested during deployment
FIXED_TOP_K = 4
DEFAULT_DOCS_DIR = "data/default_docs"
NICHE_NAME = "Legal"
NICHE_ICON = "⚖️"


def retrieve_context(query: str, top_k: int, collection_name: str) -> Tuple[List[str], List[str]]:
    # Default collection is persistent; session collections are ephemeral
    if collection_name == DEFAULT_COLLECTION:
        collection = get_or_create_collection(name=collection_name)
    else:
        collection = get_or_create_session_collection(name=collection_name)
    results = collection.query(query_texts=[query], n_results=top_k, include=["documents", "metadatas"])
    docs = results.get("documents", [[]])[0]
    metas = results.get("metadatas", [[]])[0]
    labels = [citation_label(m, i) for i, m in enumerate(metas)]
    try:
        logger.info(f"retrieve_context: collection={collection_name}, top_k={top_k}, returned={len(docs)}")
    except Exception:
        pass
    return docs, labels


 


def ensure_default_ingested() -> int:
    """Ensure the default collection has content by ingesting files from DEFAULT_DOCS_DIR.

    Returns the number of existing or newly added documents (chunk count approximation).
    """
    folder = DEFAULT_DOCS_DIR
    os.makedirs(folder, exist_ok=True)

    collection = get_or_create_collection(name=DEFAULT_COLLECTION)
    try:
        existing = collection.count()
    except Exception:
        existing = 0
    if existing > 0:
        try:
            logger.info(f"ensure_default_ingested: collection already populated, count={existing}")
        except Exception:
            pass
        return existing

    # Collect supported files
    try:
        filenames = [
            f for f in os.listdir(folder)
            if f.lower().endswith((".pdf", ".txt", ".md", ".docx"))
        ]
    except FileNotFoundError:
        filenames = []
    if not filenames:
        return 0

    chunk_size, overlap, _ = get_chunking()
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=overlap,
        separators=["\n\n", "\n", " ", ""],
    )

    all_texts: List[str] = []
    all_metas: List[dict] = []

    for name in filenames:
        path = os.path.join(folder, name)
        suffix = name.split(".")[-1].lower()
        try:
            if suffix == "pdf":
                loader = PyPDFLoader(path)
                docs = loader.load()
            elif suffix in ("txt", "md"):
                with open(path, "r", encoding="utf-8", errors="ignore") as fh:
                    content = fh.read()
                base_md = {"source": name}
                base_md.update(infer_file_metadata(name))
                docs = [LCDocument(page_content=content, metadata=base_md)]
            elif suffix == "docx":
                try:
                    import docx  # python-docx
                    doc = docx.Document(path)
                    content = "\n".join([p.text for p in doc.paragraphs])
                    base_md = {"source": name}
                    base_md.update(infer_file_metadata(name))
                    docs = [LCDocument(page_content=content, metadata=base_md)]
                except Exception:
                    # Skip DOCX if parsing fails
                    try:
                        logger.warning(f"ensure_default_ingested: failed to parse DOCX {name}; skipping")
                    except Exception:
                        pass
                    continue
            else:
                continue
        except Exception:
            try:
                logger.warning(f"ensure_default_ingested: loader error for {name}; skipping")
            except Exception:
                pass
            continue

        split_docs = splitter.split_documents(docs)
        texts = [d.page_content for d in split_docs]
        metas = []
        for i, d in enumerate(split_docs):
            md = dict(d.metadata or {})
            md.update({"source": name, "chunk_index": i})
            md.update(infer_file_metadata(name))
            md.update(extract_chunk_section(d.page_content))
            # Sanitize: Chroma metadata values must be str/int/float/bool
            clean_md = {k: v for k, v in md.items() if isinstance(v, (str, int, float, bool))}
            metas.append(clean_md)
        all_texts.extend(texts)
        all_metas.extend(metas)

    if all_texts:
        ids = [f"{all_metas[i]['source']}-{i}" for i in range(len(all_texts))]
        collection.add(ids=ids, documents=all_texts, metadatas=all_metas)
        try:
            logger.info(f"ensure_default_ingested: added_chunks={len(all_texts)} files={len(filenames)}")
        except Exception:
            pass
        return len(all_texts)

    return 0


def chat_completion(client: OpenAI, model: str, system_prompt: str, user_prompt: str) -> str:
    resp = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.2,
    )
    return resp.choices[0].message.content or ""


def main():
    st.set_page_config(page_title=f"{NICHE_NAME} RAG Chatbot", page_icon=NICHE_ICON)
    try:
        logger.info("main: render start")
    except Exception:
        pass
    # Title set dynamically below based on mode
    # Note: Streamlit's upload limit is configured at the server level (config/CLI/env),
    # not at runtime inside the script. This app relies on whatever the server allows
    # by default or as configured externally; no app-level size enforcement is applied.

    embedding_model, chat_model = get_models()

    if "messages_by_mode" not in st.session_state:
        st.session_state.messages_by_mode = {"default": [], "uploaded": []}
    if "active_collection" not in st.session_state:
        st.session_state.active_collection = DEFAULT_COLLECTION
    if "mode" not in st.session_state:
        st.session_state.mode = "default"  # or "uploaded"

    with st.sidebar:
        mode = st.radio(
            "Change mode",
            options=["Legal Chatbot", "Custom Chatbot"],
            index=0 if st.session_state.mode == "default" else 1,
            help="Legal uses pre-ingested docs. Custom uses your uploaded docs.",
        )
        if mode == "Legal Chatbot":
            st.session_state.mode = "default"
            # Clear any prior session collection when switching back
            if isinstance(st.session_state.get("active_collection"), str) and st.session_state.active_collection != DEFAULT_COLLECTION:
                try:
                    delete_session_collection(st.session_state.active_collection)
                except Exception:
                    pass
            st.session_state.active_collection = DEFAULT_COLLECTION
            with st.spinner("Preparing default legal documents…"):
                try:
                    try:
                        logger.info("ensure_default_ingested: invoked")
                    except Exception:
                        pass
                    _ = ensure_default_ingested()
                except Exception as e:
                    st.error(f"Failed to prepare default documents: {e}")
                else:
                    st.info("Using default pre-ingested documents.")
        else:
            st.session_state.mode = "uploaded"
            st.caption("Accepted: PDF, TXT, MD, DOCX. Upload size is limited by the Streamlit server settings.")
            try:
                logger.info("mode switched: uploaded")
            except Exception:
                pass
            uploaded_files = st.file_uploader(
                "Upload documents",
                type=["pdf", "txt", "md", "docx"],
                accept_multiple_files=True,
                help="Subject to the Streamlit server's configured upload limit.",
            )
            valid_files = list(uploaded_files) if uploaded_files else []

            if valid_files and st.button("Index uploaded documents", type="primary"):
                with st.spinner("Indexing documents…"):
                    try:
                        chunk_size, overlap, _ = get_chunking()
                        splitter = RecursiveCharacterTextSplitter(
                            chunk_size=chunk_size,
                            chunk_overlap=overlap,
                            separators=["\n\n", "\n", " ", ""],
                        )
                        all_texts = []
                        all_metas = []
                        for f in valid_files:
                            name = f.name
                            suffix = name.split(".")[-1].lower()
                            if suffix == "pdf":
                                tmp_path = f".tmp_upload_{name}"
                                with open(tmp_path, "wb") as out:
                                    out.write(f.getbuffer())
                                loader = PyPDFLoader(tmp_path)
                                docs = loader.load()
                            elif suffix in ("txt", "md"):
                                content = f.getvalue().decode(errors="ignore")
                                base_md = {"source": name}
                                base_md.update(infer_file_metadata(name))
                                docs = [LCDocument(page_content=content, metadata=base_md)]
                            elif suffix == "docx":
                                try:
                                    import docx  # python-docx
                                    tmp_path = f".tmp_upload_{name}"
                                    with open(tmp_path, "wb") as out:
                                        out.write(f.getbuffer())
                                    doc = docx.Document(tmp_path)
                                    content = "\n".join([p.text for p in doc.paragraphs])
                                    base_md = {"source": name}
                                    base_md.update(infer_file_metadata(name))
                                    docs = [LCDocument(page_content=content, metadata=base_md)]
                                except Exception as e:
                                    st.error(f"{name}: failed to read DOCX ({e}); skipping.")
                                    try:
                                        logger.warning(f"upload_index: failed to parse DOCX {name}; skipping")
                                    except Exception:
                                        pass
                                    continue
                            else:
                                continue

                            split_docs = splitter.split_documents(docs)
                            texts = [d.page_content for d in split_docs]
                            metas = []
                            for i, d in enumerate(split_docs):
                                md = dict(d.metadata or {})
                                md.update({"source": name, "chunk_index": i})
                                md.update(infer_file_metadata(name))
                                md.update(extract_chunk_section(d.page_content))
                                clean_md = {k: v for k, v in md.items() if isinstance(v, (str, int, float, bool))}
                                metas.append(clean_md)
                            all_texts.extend(texts)
                            all_metas.extend(metas)

                        if all_texts:
                            session_collection = f"rag_session_{hash((len(all_texts), os.urandom(4))) & 0xfffffff}"
                            collection = get_or_create_session_collection(name=session_collection)
                            ids = [f"{all_metas[i]['source']}-{i}" for i in range(len(all_texts))]
                            collection.add(ids=ids, documents=all_texts, metadatas=all_metas)
                            st.session_state.active_collection = session_collection
                            # Reset only the uploaded-mode chat when new docs are indexed
                            st.session_state.messages_by_mode["uploaded"] = []
                            st.success("Documents indexed. You can now ask questions about them.")
                            st.balloons()
                            try:
                                logger.info(f"upload_index: added_chunks={len(all_texts)} files={len(valid_files)} collection={session_collection}")
                            except Exception:
                                pass
                        else:
                            st.warning("No extractable text found in uploaded files.")
                    except Exception as e:
                        st.error(f"Indexing failed: {e}")

    # Dynamic title and personalized statement
    if st.session_state.mode == "default":
        st.title(f"{NICHE_ICON} {NICHE_NAME} RAG Chatbot")
    else:
        st.title("🧠 RAG Chatbot")

    repo_url = os.getenv("REPO_URL")
    docs_link = f"{repo_url}/tree/main/{DEFAULT_DOCS_DIR}" if repo_url else None
    if st.session_state.mode == "default":
        if docs_link:
            st.caption(
                f"I am a {NICHE_NAME} chatbot and I answer questions based on these documents: "
                f"[{DEFAULT_DOCS_DIR}]({docs_link})."
            )
        else:
            st.caption(
                f"I am a {NICHE_NAME} chatbot and I answer questions based on the pre‑ingested documents in "
                f"'{DEFAULT_DOCS_DIR}'. Set REPO_URL in your environment to show a repository link here."
            )
    else:
        pass

    # Render messages for the active mode
    active_messages = st.session_state.messages_by_mode.get(st.session_state.mode, [])
    for msg in active_messages:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])

    placeholder = (
        "Ask a legal/compliance question grounded in the default documents…"
        if st.session_state.mode == "default"
        else "Ask a question about your uploaded documents…"
    )
    user_input = st.chat_input(placeholder)
    if user_input:
        st.session_state.messages_by_mode[st.session_state.mode].append({"role": "user", "content": user_input})
        with st.chat_message("user"):
            st.markdown(user_input)

        # Ensure API key configured (robust to BOM issues on Windows)
        try:
            openrouter_api_key = get_env("OPENROUTER_API_KEY")
        except Exception:
            st.error("Missing OPENROUTER_API_KEY. Set it in environment or .env.")
            return

        try:
            _, _, cfg_top_k = get_chunking()
            top_k = max(cfg_top_k, 6)
            docs, labels = retrieve_context(user_input, top_k, st.session_state.active_collection)
            try:
                logger.info("retrieval_preview: query='%s' hits=%d labels=%s", user_input, len(docs), labels)
            except Exception:
                pass
        except Exception as e:
            st.error(f"Retrieval failed: {e}")
            return

        # If no context, return concise fallback with no disclaimer or safety
        if not docs:
            fallback = (
                "The data I am trained on doesn't provide an answer to that question."
                if st.session_state.mode == "default" else "I don't know."
            )
            with st.chat_message("assistant"):
                st.markdown(fallback)
            st.session_state.messages_by_mode[st.session_state.mode].append({"role": "assistant", "content": fallback})
            return

        context_block = format_context(docs, labels)
        repo_url = os.getenv("REPO_URL")
        if repo_url and not repo_url.startswith(("http://", "https://")):
            repo_url = f"https://{repo_url}"
        system_prompt = build_system_prompt(mode=st.session_state.mode, repo_url=repo_url, docs_dir=DEFAULT_DOCS_DIR)
        user_prompt = build_user_prompt(user_input, context_block, mode=st.session_state.mode, repo_url=repo_url, docs_dir=DEFAULT_DOCS_DIR)

        # LangChain chat model with OpenRouter backend
        llm = ChatOpenAI(
            model=chat_model,
            temperature=0.2,
            api_key=openrouter_api_key,
            base_url=os.getenv("OPENROUTER_BASE_URL", "https://openrouter.ai/api/v1"),
        )
        try:
            with st.spinner("Thinking…"):
                draft = llm.invoke([SystemMessage(content=system_prompt), HumanMessage(content=user_prompt)]).content
        except Exception as e:
            st.error("Connection issue reaching the model provider. Please check your internet or try again.")
            try:
                logger.info("chat_invoke_error: %s", e)
            except Exception:
                pass
            return

        # Enforce: when context exists, do not allow fallback phrasing in the draft
        fallback_default = "The data I am trained on doesn't provide an answer to that question."
        fallback_custom = "I don't know."
        for bad in (fallback_default, fallback_custom):
            if bad in (draft or ""):
                draft = (draft or "").replace(bad, "").strip()

        with st.chat_message("assistant"):
            st.markdown(draft)
            st.session_state.messages_by_mode[st.session_state.mode].append({"role": "assistant", "content": draft})


if __name__ == "__main__":
    main()


